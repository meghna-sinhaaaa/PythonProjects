import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

#source_doc = Document('mainone.docx')
inputt = input("Enter main file name: ")
source_doc = Document(inputt)
source_header = source_doc.sections[0].header

p = source_header.paragraphs

text = ""
for para in p:
    print(para.text)
    text = text+para.text

for filename in os.listdir():
    if filename.endswith('.docx') and filename != 'mainone.docx':


        print(filename)
        target_doc = Document(filename)
        target_header = target_doc.sections[0].header
        htable = target_header.add_table(1,2,Inches(6))
        htab_cells = htable.rows[0].cells
        ht0 = htab_cells[0].add_paragraph()
        kh = ht0.add_run()
        kh.add_picture("cherries.png", width=Inches(1))
        ht1 = htab_cells[1].add_paragraph(text)
        ht1.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        target_doc.save(filename)


